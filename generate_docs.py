import docx
from docx.shared import Pt, Inches

def add_header(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_prompt(doc, title, prompt_text):
    doc.add_heading(title, level=3)
    p = doc.add_paragraph()
    p.add_run(prompt_text).font.name = 'Courier New'

def generate_technical_docs():
    doc = docx.Document()
    add_header(doc, 'Detailed Technical Documentation - VitalSync v2.0', 1)
    
    add_header(doc, '1. Project Environment Setup', 2)
    doc.add_paragraph("Ensure you have Node.js v18+ and a standard IDE like VS Code.")
    doc.add_paragraph("Installation steps:\n1. Clone the repository\n2. Run `npm install`\n3. Populate `.env` with Firebase configuration\n4. Run `npm run dev` to start Vite.")
    
    add_header(doc, '2. Core Directory Structure', 2)
    doc.add_paragraph("src/api/ - Contains firestore.js, acting as the interface boundary between React and Firebase SDK. Handles CRUD operations.")
    doc.add_paragraph("src/components/ - Reusable UI widgets. Contains subfolders for specialized areas (e.g., charts/).")
    doc.add_paragraph("src/context/ - Global state providers. AuthContext wraps the app to provide user identity, while ThemeContext handles Dark/Light mode preferences.")
    doc.add_paragraph("src/pages/ - High-level route components mapped to React Router v6 endpoints.")
    
    add_header(doc, '3. Data Flow & State Management', 2)
    doc.add_paragraph("Authentication state is globally provided via AuthContext. This avoids prop drilling. Real-time auth listeners (onAuthStateChanged) keep the state synchronized.")
    doc.add_paragraph("Page-level data (e.g., fetching today's logs) is handled within individual page components (Dashboard.jsx) using the useEffect hook and local useState.")
    
    add_header(doc, '4. Styling System', 2)
    doc.add_paragraph("The application leverages Tailwind CSS heavily. Custom configuration allows the use of arbitrary values and complex dark mode selectors.")
    doc.add_paragraph("Dark mode is enabled by adding the 'dark' class to the HTML root node. Tailwind's 'dark:' variant handles all localized styling changes.")
    
    add_header(doc, '5. Deployment Strategy', 2)
    doc.add_paragraph("Vercel is the recommended hosting provider. Since this is a Vite-based SPA, the build command is `npm run build` and output directory is `dist`.")
    doc.add_paragraph("Environment variables must be manually mapped in Vercel's project settings to mimic the `.env` local file.")
    
    doc.save('docs/TECHNICAL_DOCS.docx')


def generate_prompt_docs():
    doc = docx.Document()
    add_header(doc, 'VitalSync - Vibecoding Prompts Reference', 1)
    
    doc.add_paragraph("This document contains the exact sequence of prompts utilized to construct the VitalSync application from the Antigravity Build Guide.")
    
    add_header(doc, 'Phase 1 - Scaffolding', 2)
    add_prompt(doc, 'Project Initialization', "Create a new React.js project called 'vitalsync' using Vite. Set up Tailwind CSS v3. Install firebase, react-router-dom, d3, react-icons, date-fns. Create the exact folder structure inside src/ including api, components, context, firebase, and pages. Setup App.jsx with React Router v6.")
    
    add_header(doc, 'Phase 2 - Firebase Setup', 2)
    add_prompt(doc, '2A: Firebase Config', "In src/firebase/config.js, initialize Firebase using the config structure... export auth and db.")
    add_prompt(doc, '2B: Auth Context', "Build a complete AuthContext in src/context/AuthContext.jsx using Firebase Auth. Expose user, loading, isAdmin, login, register, logout.")
    add_prompt(doc, '2C: Protected Routes', "Create ProtectedRoute.jsx and AdminRoute.jsx. In App.jsx, wrap the dashboard routes with ProtectedRoute and /admin with AdminRoute.")
    
    add_header(doc, 'Phase 3 - Database API', 2)
    add_prompt(doc, 'Firestore Wrapper', "Create src/api/firestore.js with all Firestore helper functions. Write addHealthLog, getHealthLogs, getHealthLogByDate, updateHealthLog, deleteHealthLog, addGoal, getGoals, updateGoal, deleteGoal, getAllUsers, getAllHealthLogs.")
    
    add_header(doc, 'Phase 4 - Core CRUD', 2)
    add_prompt(doc, '4A: Log Form', "Build src/components/LogForm.jsx — a form to log daily health metrics with fields for Date, Steps, Heart Rate, Sleep, Hydration, Weight. On mount, check if log exists and pre-fill.")
    add_prompt(doc, '4B: Health Log Page', "Build src/pages/HealthLog.jsx with two sections: Today's Log (LogForm) and Recent History table.")
    
    add_header(doc, 'Phase 5 - D3.js Charts', 2)
    add_prompt(doc, '5A: Steps Chart', "Build StepsChart.jsx using D3.js. Horizontal bar chart showing last 14 days.")
    add_prompt(doc, '5B: Heart Rate Chart', "Build HeartRateChart.jsx using D3.js. Smooth line chart with normal range highlighting.")
    add_prompt(doc, '5C: Sleep & Hydration', "Build SleepChart.jsx (Area chart) and HydrationChart.jsx (Vertical bar chart).")
    add_prompt(doc, '5D: Analytics Page', "Build Analytics.jsx fetching logs and displaying 4 metric cards and 4 charts stacked vertically.")
    
    add_header(doc, 'Phase 6 - Goals', 2)
    add_prompt(doc, '6A: Goal Card', "Build GoalCard.jsx with title, circular progress ring (SVG), deadline, and status badge.")
    add_prompt(doc, '6B: Goals Page', "Build Goals.jsx with 'Add New Goal' form and 'My Goals' list broken into active vs completed.")
    
    add_header(doc, 'Phase 7 - Admin', 2)
    add_prompt(doc, 'Admin Dashboard', "Build AdminDashboard.jsx with Platform Overview stats, All Users Table with expandable logs, and a System Health Metrics chart.")
    
    add_header(doc, 'Phase 8 - Notifications', 2)
    add_prompt(doc, 'Notifications System', "Build an in-app notification system. Create a bell icon in Navbar, create notifications collection, auto-generate notifications on streak or low activity.")
    
    add_header(doc, 'Phase 9 - Landing & Auth', 2)
    add_prompt(doc, '9A: Landing Page', "Build Landing.jsx with Hero section, Features section, and How It Works.")
    add_prompt(doc, '9B: Auth Pages', "Build Register.jsx and Login.jsx with form validation, error handling, and redirection.")
    
    add_header(doc, 'Phase 10 - Dashboard', 2)
    add_prompt(doc, 'User Dashboard', "Build Dashboard.jsx. Show Today's Summary (4 metric cards), Quick Log Button, Weekly Progress mini charts, and Active Goals.")
    
    add_header(doc, 'Phase 11 - Navbar', 2)
    add_prompt(doc, 'Navbar Component', "Build Navbar.jsx with logo, navigation links, notification bell, and user dropdown.")
    
    doc.save('docs/PROMPT_DOCS.docx')

if __name__ == '__main__':
    generate_technical_docs()
    generate_prompt_docs()
    print("Generated DOCX files successfully.")
